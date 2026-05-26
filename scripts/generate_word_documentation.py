"""
Generiše Word dokument sa objašnjenjem koda i genomike za projekat.
Izlaz: deliverables/Project_Documentation.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Single-Cell Analysis of Immune Response\n"
        "to Nanoplastic Particles\n\n"
        "Code & Genomic Informatics Guide"
    )
    r.bold = True
    r.font.size = Pt(20)
    doc.add_paragraph("Master project – Genomic Informatics").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- PART 1: BIOLOGY & GENOMICS ---
    add_heading(doc, "Part 1: Biological and Genomic Background", 1)

    add_heading(doc, "1.1 What is the biological question?", 2)
    doc.add_paragraph(
        "Nanoplastics (tiny plastic particles, often < 1 µm) can enter the bloodstream and interact "
        "directly with immune cells. Particle size (e.g. 40 nm vs 200 nm) may change how strongly "
        "and in which cell types the immune system responds."
    )
    doc.add_paragraph(
        "This project uses single-cell RNA sequencing (scRNA-seq) on human peripheral blood "
        "immune cells (PBMC) from one donor exposed to:"
    )
    add_bullets(
        doc,
        [
            "Sample 1: 40 nm carboxylated polystyrene nanoparticles (PSNPs)",
            "Sample 2: 200 nm PSNPs",
            "Sample 3: mixture of 40 nm + 200 nm",
            "Sample 4: control (no exposure)",
        ],
    )

    add_heading(doc, "1.2 What is scRNA-seq?", 2)
    doc.add_paragraph(
        "Bulk RNA-seq measures average gene expression across millions of cells mixed together. "
        "scRNA-seq captures RNA from thousands of individual cells separately."
    )
    doc.add_paragraph("Key concepts:")
    add_bullets(
        doc,
        [
            "Gene expression: how much each gene is transcribed into mRNA in a cell.",
            "UMI / counts matrix: table of genes (rows) × cells (columns) with read counts.",
            "Cell barcode: unique ID for each captured cell (e.g. AAACCCAAGACGCATG-1).",
            "Feature: usually a gene; sometimes other features in the assay.",
        ],
    )

    add_heading(doc, "1.3 PBMC and immune cell types", 2)
    doc.add_paragraph(
        "PBMC (peripheral blood mononuclear cells) include lymphocytes and monocytes. "
        "Major types in this project:"
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Cell type"
    hdr[1].text = "Role"
    hdr[2].text = "Example markers"
    rows = [
        ("T cells (CD4+)", "Adaptive immunity, help/coordinate", "CD3D, CD3E, IL7R"),
        ("Cytotoxic T / CD8+", "Kill infected/abnormal cells", "NKG7, GNLY, PRF1, GZMB"),
        ("B cells", "Antibody production", "MS4A1, CD79A, CD74"),
        ("NK cells", "Innate killing", "NKG7, KLRD1"),
        ("Monocytes (CD14+)", "Phagocytosis, inflammation", "LYZ, S100A8, S100A9"),
        ("Dendritic cells", "Antigen presentation", "FCER1A, CST3"),
    ]
    for ct, role, markers in rows:
        row = table.add_row().cells
        row[0].text = ct
        row[1].text = role
        row[2].text = markers

    add_heading(doc, "1.4 Core bioinformatics concepts used in this pipeline", 2)

    concepts = [
        (
            "Quality control (QC)",
            "Remove damaged cells, empty droplets, and doublets before analysis. "
            "Metrics: number of genes detected, total UMI counts, % mitochondrial genes (MT-). "
            "High MT% often means dying/stressed cells.",
        ),
        (
            "Normalization",
            "Cells have different sequencing depth. normalize_total scales each cell to the same "
            "total counts; log1p stabilizes variance for downstream statistics.",
        ),
        (
            "Highly variable genes (HVG)",
            "Genes that vary most across cells carry biological signal; ~3000 HVGs are used "
            "for dimensionality reduction instead of all ~20,000 genes.",
        ),
        (
            "PCA (Principal Component Analysis)",
            "Linear reduction: thousands of genes → ~30 principal components capturing main variation.",
        ),
        (
            "Batch correction / integration",
            "Four samples were processed separately → technical 'batch' effects. Combat adjusts "
            "expression in PCA space so cells cluster by biology, not by sample file.",
        ),
        (
            "UMAP",
            "Non-linear 2D embedding for visualization. Similar cells appear close together.",
        ),
        (
            "Leiden clustering",
            "Groups cells into clusters (communities) on a nearest-neighbor graph.",
        ),
        (
            "Cell type annotation",
            "Assigning labels (T cell, B cell, …) using marker genes or reference atlases (Azimuth, CoDi).",
        ),
        (
            "Differential expression (DE)",
            "Which genes are up/down in exposed vs control, within one cell type. "
            "Wilcoxon rank-sum test; adjusted p-value (FDR) controls false discoveries.",
        ),
        (
            "Pathway enrichment",
            "Are DE genes over-represented in GO, KEGG, or Reactome pathways? "
            "Links gene lists to biology (e.g. interferon response, inflammation).",
        ),
        (
            "Size-specific effects",
            "Comparing DE gene sets: unique to 40 nm, unique to 200 nm, shared, or only in mixture.",
        ),
    ]
    for name, desc in concepts:
        p = doc.add_paragraph()
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # --- PART 2: PROJECT STRUCTURE ---
    add_heading(doc, "Part 2: Project Structure and Workflow", 1)

    add_heading(doc, "2.1 Folder layout", 2)
    add_bullets(
        doc,
        [
            "config/config.yaml – all parameters (QC, DE, markers, paths)",
            "data/raw/ – downloaded .h5ad and CoDi CSV (not on GitHub)",
            "data/processed/ – integrated_annotated.h5ad after pipeline",
            "results/figures/ – UMAP plots, composition barplot",
            "results/tables/ – CSV outputs (DE, pathways, scores)",
            "scripts/ – Python and R code",
            "deliverables/ – PowerPoint, this Word doc, video script",
        ],
    )

    add_heading(doc, "2.2 How to run (reproducibility)", 2)
    add_bullets(
        doc,
        [
            "pip install -r requirements.txt",
            "python scripts/download_data.py",
            "python scripts/run_pipeline.py",
            "python scripts/make_slides.py",
            "(optional) Rscript scripts/azimuth_annotation.R",
        ],
    )

    add_heading(doc, "2.3 Data format: AnnData (.h5ad)", 2)
    doc.add_paragraph(
        "AnnData is the standard Python object for scRNA-seq. It contains:"
    )
    add_bullets(
        doc,
        [
            "adata.X – expression matrix (cells × genes)",
            "adata.obs – metadata per cell (condition, cell type, QC metrics)",
            "adata.var – metadata per gene",
            "adata.layers – e.g. raw counts before normalization",
            "adata.obsm – embeddings (PCA, UMAP coordinates)",
        ],
    )

    doc.add_page_break()

    # --- PART 3: CODE WALKTHROUGH ---
    add_heading(doc, "Part 3: Code Explanation (function by function)", 1)

    functions = [
        (
            "scripts/download_data.py",
            [
                ("load_config()", "Reads Zenodo record ID and paths from config.yaml."),
                (
                    "download_file(url, path)",
                    "Downloads one file with streaming; skips if already present.",
                ),
                (
                    "main()",
                    "Calls Zenodo API (record 15866724), saves metadata JSON, downloads "
                    "all .h5ad, .csv, and .rds files into data/raw/.",
                ),
            ],
        ),
        (
            "scripts/run_pipeline.py – main analysis",
            [
                ("load_config()", "Loads YAML parameters."),
                ("setup_paths()", "Creates data/ and results/ folders."),
                (
                    "read_and_qc_sample()",
                    "Loads one .h5ad; adds sample_id and condition; computes MT%; "
                    "filters cells: min_genes≥200, max_genes≤7000, min_counts≥500, MT%≤15.",
                ),
                (
                    "merge_and_integrate()",
                    "Concatenates 4 samples; normalizes; selects 3000 HVGs; scales; PCA; "
                    "Combat batch correction; builds neighbor graph; UMAP; Leiden clusters.",
                ),
                (
                    "marker_based_annotation()",
                    "For each cell, scores mean expression of marker sets; assigns cell type "
                    "with highest score → adata.obs['cell_type_marker'].",
                ),
                (
                    "load_codi_annotations()",
                    "Maps CoDi labels from *_CoDi_KLD.csv to cells for validation.",
                ),
                (
                    "composition_analysis()",
                    "Counts cell types per condition; saves CSV and barplot.",
                ),
                (
                    "differential_expression_by_celltype()",
                    "For each cell type and each exposure (40nm, 200nm, mix): Wilcoxon test "
                    "vs control → differential_expression_all.csv.",
                ),
                (
                    "pathway_enrichment()",
                    "Enrichr via gseapy: GO, KEGG, Reactome on significant upregulated genes.",
                ),
                (
                    "size_specific_effects()",
                    "Classifies DE genes: unique_40nm, unique_200nm, shared_40_200, "
                    "shared_all_three, mix_only_emergent.",
                ),
                (
                    "additional_insights()",
                    "Five extra analyses: cell cycle, IFN score, antigen presentation, "
                    "pseudobulk matrix, CoDi vs marker agreement.",
                ),
                (
                    "save_core_figures()",
                    "Exports UMAP colored by condition, cluster, and cell type.",
                ),
                (
                    "main()",
                    "Runs all steps in order; writes integrated_annotated.h5ad.",
                ),
            ],
        ),
        (
            "scripts/azimuth_annotation.R",
            [
                (
                    "RunAzimuth()",
                    "Maps cells to Satija PBMC reference; outputs predicted cell types "
                    "at three resolution levels + confidence score.",
                ),
            ],
        ),
        (
            "scripts/make_slides.py",
            [
                ("main()", "Builds PowerPoint from results figures and size-specific summary table."),
            ],
        ),
    ]

    for script_name, funcs in functions:
        add_heading(doc, script_name, 2)
        for fname, fdesc in funcs:
            p = doc.add_paragraph()
            p.add_run(fname).bold = True
            p.add_run(" – " + fdesc)

    doc.add_page_break()

    # --- PART 4: CONFIG ---
    add_heading(doc, "Part 4: Configuration (config/config.yaml)", 1)
    doc.add_paragraph("Key parameters and why they were chosen:")
    cfg_items = [
        ("qc.min_genes = 200", "Standard lower bound for real PBMC cells."),
        ("qc.max_genes = 7000", "Filters likely doublets (two cells in one droplet)."),
        ("qc.max_mt_percent = 15", "Removes stressed/dying cells."),
        ("preprocessing.n_hvgs = 3000", "Common choice for ~10k–50k cells."),
        ("preprocessing.integration_method = combat", "Works on Windows; corrects batch between 4 samples."),
        ("de.pval_adj_threshold = 0.05", "Standard FDR cutoff."),
        ("de.logfc_threshold = 0.25", "Minimum biological effect size on log scale."),
        ("de.min_cells_per_group = 20", "Minimum cells needed for stable DE test."),
    ]
    for k, v in cfg_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(k + ": ").bold = True
        p.add_run(v)

    doc.add_page_break()

    # --- PART 5: OUTPUTS ---
    add_heading(doc, "Part 5: Output Files and How to Interpret Them", 1)

    outputs = [
        (
            "umap_condition.png",
            "Each point = one cell. Colors = treatment. Overlap of colors in same region "
            "= similar expression despite different exposure.",
        ),
        (
            "umap_celltypes_marker.png",
            "Shows immune compartments (T, B, monocytes). Exposure effects should be "
            "interpreted within each type.",
        ),
        (
            "composition_barplot.png",
            "Fraction of each cell type per condition. Shifts suggest nanoplastics change "
            "immune cell proportions (e.g. more monocytes).",
        ),
        (
            "differential_expression_all.csv",
            "Columns: gene names, logFC, pvals, pvals_adj, cell_type, comparison. "
            "Positive logFC = higher in exposure than control.",
        ),
        (
            "pathway_enrichment_all.csv",
            "Enriched GO/KEGG/Reactome terms for DE genes – biological interpretation.",
        ),
        (
            "size_specific_effects_summary.csv",
            "Counts of DE genes in each size category per cell type.",
        ),
        (
            "ifn_scores_by_condition.csv",
            "Mean interferon signature per condition – innate immune activation.",
        ),
    ]
    for fname, meaning in outputs:
        p = doc.add_paragraph()
        p.add_run(fname + ": ").bold = True
        p.add_run(meaning)

    doc.add_page_break()

    # --- PART 6: PRESENTATION ---
    add_heading(doc, "Part 6: How to Explain Results (for defense / report)", 1)

    add_heading(doc, "6.1 Suggested narrative", 2)
    steps = [
        "Introduce nanoplastic health concern and size-dependent hypothesis.",
        "Describe experimental design: 4 conditions, one donor PBMC, scRNA-seq.",
        "Show QC: how many cells removed and why thresholds are justified.",
        "Show integrated UMAP: batch correction worked; biology visible.",
        "Show cell type annotation and composition changes.",
        "Present DE and pathways for 1–2 most affected cell types.",
        "Highlight size-specific gene sets and biological interpretation.",
        "Mention limitations: single donor, in vitro exposure, need validation.",
    ]
    add_bullets(doc, steps)

    add_heading(doc, "6.2 Glossary (quick reference)", 2)
    glossary = [
        ("AnnData", "Annotated data matrix for single-cell experiments in Python."),
        ("Combat", "Batch correction method removing systematic sample effects."),
        ("CoDi", "Reference-based cell type labels provided with the dataset."),
        ("DE", "Differential expression – statistical comparison between groups."),
        ("FDR / padj", "False discovery rate – corrected p-value for multiple testing."),
        ("HVG", "Highly variable genes."),
        ("logFC", "Log2 fold change between two groups."),
        ("PBMC", "Peripheral blood mononuclear cells."),
        ("PSNP", "Polystyrene nanoparticle."),
        ("Pseudobulk", "Summing counts per cell type and condition (bulk-like profile)."),
        ("Scanpy", "Python library for scRNA-seq analysis."),
        ("UMI", "Unique molecular identifier – deduplicated read count."),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph()
        p.add_run(term + " – ").bold = True
        p.add_run(defn)

    return doc


def main():
    out_dir = Path("deliverables")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Project_Documentation.docx"

    doc = build_document()
    doc.save(out_path)
    print(f"Saved: {out_path.resolve()}")


if __name__ == "__main__":
    main()
